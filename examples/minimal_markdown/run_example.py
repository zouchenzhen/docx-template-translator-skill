"""Run the minimal_markdown example end-to-end.

Steps:
  1. build_template.py            -> sample_template.docx   (skipped if --template is given)
  2. inspect_docx_template.py     -> sample_template.template-report.json
  3. pandoc (if available) OR a python-docx fallback   -> body.docx
  4. adaptive_docx_pipeline.py    -> final.docx
  5. (Windows + Word only) finalize_word_docx.py + render_pdf_preview.py

The script is meant to be runnable from anywhere — it resolves all paths
relative to itself, not to the current working directory.

Optional: pass --template <path-to-real.docx> to run the example against your
own real Word template (e.g. a school thesis template). The given file is
copied into examples/minimal_markdown/sample_template.docx so the rest of the
pipeline keeps working unchanged. Without --template the example regenerates
the lightweight reproducible sample template via build_template.py.
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parent.parent
SCRIPTS = REPO_ROOT / "skills" / "docx-template-translator" / "scripts"
PRESETS = REPO_ROOT / "skills" / "docx-template-translator" / "presets"

TEMPLATE = HERE / "sample_template.docx"
SAMPLE_MD = HERE / "sample.md"
BODY = HERE / "body.docx"
FINAL = HERE / "final.docx"
REPORT = HERE / "sample_template.template-report.json"


def run(cmd: list[str]) -> None:
    print(">>", " ".join(str(c) for c in cmd))
    subprocess.run(cmd, check=True)


def step_build_template() -> None:
    run([sys.executable, str(HERE / "build_template.py")])


def step_inspect() -> None:
    run([
        sys.executable,
        str(SCRIPTS / "inspect_docx_template.py"),
        str(TEMPLATE),
        "--out",
        str(REPORT),
    ])


def step_rough_body_with_pandoc() -> bool:
    if shutil.which("pandoc") is None:
        return False
    try:
        run([
            "pandoc",
            str(SAMPLE_MD),
            "--resource-path",
            str(HERE),
            "--reference-doc",
            str(TEMPLATE),
            "-o",
            str(BODY),
        ])
        return True
    except subprocess.CalledProcessError:
        return False


def step_rough_body_fallback() -> None:
    """python-docx-only fallback: build a minimal body.docx without pandoc.

    Only invoked when pandoc isn't on PATH so the example still completes in
    minimal environments. Real users should install pandoc for proper
    Markdown -> DOCX conversion.
    """
    from docx import Document

    doc = Document()
    text = SAMPLE_MD.read_text(encoding="utf-8")
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    doc.save(BODY)
    print(f"(fallback) wrote {BODY}")


def step_adaptive_pipeline() -> None:
    run([
        sys.executable,
        str(SCRIPTS / "adaptive_docx_pipeline.py"),
        "--template",
        str(TEMPLATE),
        "--body-docx",
        str(BODY),
        "--out",
        str(FINAL),
        "--config",
        str(PRESETS / "zhengzhou_thesis.json"),
        "--three-line-tables",
    ])


def step_finalize_if_possible(pages: str) -> None:
    if sys.platform != "win32":
        print("(skip finalize: not on Windows)")
        return
    try:
        import win32com.client  # noqa: F401
    except ImportError:
        print("(skip finalize: pywin32 not installed)")
        return
    try:
        run([
            sys.executable,
            str(SCRIPTS / "finalize_word_docx.py"),
            str(FINAL),
            "--pdf",
        ])
        run([
            sys.executable,
            str(SCRIPTS / "render_pdf_preview.py"),
            str(FINAL.with_suffix(".pdf")),
            "--pages",
            pages,
        ])
    except subprocess.CalledProcessError as exc:
        print(f"(finalize step failed: {exc})")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--template",
        default=None,
        help=(
            "Optional path to a real .docx template (e.g. a school thesis "
            "template). The file is copied into sample_template.docx so the "
            "rest of the pipeline runs unchanged. Without this flag, "
            "build_template.py regenerates the bundled minimal sample."
        ),
    )
    parser.add_argument(
        "--preview-pages",
        default="1-4",
        help=(
            "Page range passed to render_pdf_preview.py (default '1-4'). "
            "When running against a real thesis template that bundles its own "
            "front-matter + sample chapters, prefer pages that show real "
            "body content (TOC + figure/equation/table sample + references), "
            "e.g. '7,12,13,14'."
        ),
    )
    args = parser.parse_args()

    if args.template:
        custom = Path(args.template).expanduser().resolve()
        if not custom.is_file():
            print(f"--template path not found: {custom}")
            return 2
        TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(custom, TEMPLATE)
        print(f"using custom template: {custom}")
        print(f"  copied to: {TEMPLATE}")
    else:
        step_build_template()
    step_inspect()
    if not step_rough_body_with_pandoc():
        print("(pandoc unavailable — falling back to python-docx body builder)")
        step_rough_body_fallback()
    step_adaptive_pipeline()
    step_finalize_if_possible(args.preview_pages)
    print()
    print("done. outputs in", HERE)
    return 0


if __name__ == "__main__":
    for _s in (sys.stdout, sys.stderr):
        if hasattr(_s, "reconfigure"):
            _s.reconfigure(encoding="utf-8", errors="replace")

    sys.exit(main())
