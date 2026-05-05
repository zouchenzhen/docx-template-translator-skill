"""Build a minimal sample_template.docx for the example.

We do NOT commit a binary template into the repo. Instead, we ship this
generator so the example is reproducible, the template's defined styles are
visible in source, and reviewers can diff changes without binary noise.

The generated template is intentionally minimal and not tied to any specific
school. It defines a small set of styles a real institutional template would
have, including one Chinese-named body style ("论文正文") and a misleading
"Body Text" style — both deliberately echo the kind of structure where pandoc
--reference-doc alone would misroute paragraphs.

Run from anywhere; the output goes next to this script.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT = HERE / "sample_template.docx"


def _ensure_paragraph_style(doc, name: str) -> None:
    """Add a paragraph style if it's not already there."""
    if name in [s.name for s in doc.styles]:
        return
    doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def main() -> int:
    doc = Document()

    # Add a paragraph style that's named in Chinese — this is the body style
    # we'd remap into in a Chinese-thesis preset run.
    _ensure_paragraph_style(doc, "论文正文")
    body_style = doc.styles["论文正文"]
    body_style.font.name = "Times New Roman"
    body_style.font.size = Pt(12)
    body_style.paragraph_format.first_line_indent = Pt(24)

    # Add a misleading "Body Text" style (large + bold) as a stand-in for the
    # cover-page-style trap described in the case study.
    _ensure_paragraph_style(doc, "Cover Title")
    cover = doc.styles["Cover Title"]
    cover.font.name = "Times New Roman"
    cover.font.size = Pt(28)
    cover.font.bold = True
    cover.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Front matter — a tiny "cover" so reviewers see the template doesn't
    # start blank. The real adaptive pipeline appends the body after this.
    front_para = doc.add_paragraph("Sample Template — Cover")
    front_para.style = doc.styles["Cover Title"]

    abstract_heading = doc.add_paragraph("摘要")
    abstract_heading.style = doc.styles["Heading 1"]

    # Leave one body paragraph in the template body so inspect_docx_template
    # has at least one used-in-body sample to summarize.
    sample_body = doc.add_paragraph(
        "这是模板自带的一段示例正文，用来验证检查脚本能识别到 论文正文 样式。"
    )
    sample_body.style = doc.styles["论文正文"]

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
