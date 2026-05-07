#!/usr/bin/env python
"""Validate a reconstructed DOCX/PDF against common template-conversion failures.

Usage:
  python validate_docx_conversion.py final.docx --pdf final.pdf --out validation.json
  python validate_docx_conversion.py final.docx --ordered-term 绪论 --ordered-term 参考文献
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DEFAULT_PLACEHOLDERS = [
    "李四",
    "王五",
    "张三",
    "lorem ipsum",
    "Lorem ipsum",
    "为了提高本科生学位论文的质量",
    "此处为论文题目的英文翻译",
    "题目（一般不宜超过25字",
    "右键更新域或由自动最终化脚本更新目录",
]


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def paragraph_style_name(paragraph) -> str:
    try:
        return paragraph.style.name
    except Exception:
        return ""


def collect_docx_report(
    docx_path: Path,
    *,
    placeholders: list[str],
    ordered_terms: list[str],
    required_headings: list[str],
    forbidden_header_terms: list[str],
    min_images: int,
    min_tables: int,
) -> tuple[dict, list[str]]:
    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    compact_text = normalize(text)
    failures: list[str] = []

    placeholder_hits = [term for term in placeholders if term and term in text]
    if placeholder_hits:
        failures.append("template placeholder text remains")

    order_positions = {term: compact_text.find(normalize(term)) for term in ordered_terms}
    if ordered_terms:
        previous = -1
        for term in ordered_terms:
            pos = order_positions[term]
            if pos < 0:
                failures.append(f"ordered term missing: {term}")
            elif pos <= previous:
                failures.append(f"ordered term out of order: {term}")
            previous = pos

    heading_styles = {"Heading 1", "Heading 2", "Heading 3", "标题 1", "标题 2", "标题 3"}
    headings = [
        {"text": p.text.strip(), "style": paragraph_style_name(p)}
        for p in doc.paragraphs
        if paragraph_style_name(p) in heading_styles
    ]
    heading_compact = {normalize(item["text"]): item for item in headings}
    missing_heading_terms = [
        term for term in required_headings if normalize(term) not in heading_compact
    ]
    if missing_heading_terms:
        failures.append("required headings are not heading-styled")

    section_headers = []
    forbidden_header_hits = []
    for idx, section in enumerate(doc.sections):
        header_text = "\n".join(p.text for p in section.header.paragraphs).strip()
        section_headers.append({"section": idx, "text": header_text})
        for term in forbidden_header_terms:
            if term and term in header_text:
                forbidden_header_hits.append({"section": idx, "term": term, "text": header_text})
        if forbidden_header_terms:
            # Also catch stale header references that may render through Word even
            # when python-docx sees little visible text.
            refs = section._sectPr.findall(qn("w:headerReference"))
            if refs and not header_text:
                section_headers[-1]["header_reference_count"] = len(refs)
    if forbidden_header_hits:
        failures.append("forbidden back-matter header inherited by body section")

    image_count = len(doc.inline_shapes)
    table_count = len(doc.tables)
    if image_count < min_images:
        failures.append(f"image count below minimum: {image_count} < {min_images}")
    if table_count < min_tables:
        failures.append(f"table count below minimum: {table_count} < {min_tables}")

    report = {
        "docx": str(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": table_count,
        "image_count": image_count,
        "placeholder_hits": placeholder_hits,
        "ordered_terms": order_positions,
        "heading_count": len(headings),
        "missing_heading_terms": missing_heading_terms,
        "section_headers": section_headers,
        "forbidden_header_hits": forbidden_header_hits,
    }
    return report, failures


def collect_pdf_report(pdf_path: Path, placeholders: list[str]) -> tuple[dict, list[str]]:
    failures: list[str] = []
    try:
        import fitz
    except ImportError:
        return {"pdf": str(pdf_path), "error": "PyMuPDF not installed"}, ["pdf validation unavailable"]

    with fitz.open(pdf_path) as pdf:
        page_texts = [page.get_text() for page in pdf]
    text = "\n".join(page_texts)
    placeholder_hits = [term for term in placeholders if term and term in text]
    if placeholder_hits:
        failures.append("template placeholder text remains in PDF")
    report = {
        "pdf": str(pdf_path),
        "page_count": len(page_texts),
        "placeholder_hits": placeholder_hits,
    }
    return report, failures


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--pdf", default=None)
    parser.add_argument("--out", default=None)
    parser.add_argument("--placeholder", action="append", default=[])
    parser.add_argument("--no-default-placeholders", action="store_true")
    parser.add_argument("--ordered-term", action="append", default=[])
    parser.add_argument("--required-heading", action="append", default=[])
    parser.add_argument("--forbidden-header", action="append", default=[])
    parser.add_argument("--min-images", type=int, default=0)
    parser.add_argument("--min-tables", type=int, default=0)
    args = parser.parse_args()

    docx_path = Path(args.docx)
    placeholders = list(args.placeholder)
    if not args.no_default_placeholders:
        placeholders = DEFAULT_PLACEHOLDERS + placeholders

    docx_report, failures = collect_docx_report(
        docx_path,
        placeholders=placeholders,
        ordered_terms=args.ordered_term,
        required_headings=args.required_heading,
        forbidden_header_terms=args.forbidden_header,
        min_images=args.min_images,
        min_tables=args.min_tables,
    )
    report = {"docx": docx_report}

    if args.pdf:
        pdf_report, pdf_failures = collect_pdf_report(Path(args.pdf), placeholders)
        report["pdf"] = pdf_report
        failures.extend(pdf_failures)

    status = "PASS" if not failures else "FAIL"
    report["status"] = status
    report["failures"] = failures

    output = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out:
        Path(args.out).write_text(output, encoding="utf-8")
    print(f"STATUS: {status}")
    if failures:
        for failure in failures:
            print(f"- {failure}")
    if args.out:
        print(args.out)
    return 0 if status == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
