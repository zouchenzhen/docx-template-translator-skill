#!/usr/bin/env python
"""Inspect a Word template and write a compact JSON report.

Usage:
  python inspect_docx_template.py template.docx --out template_report.json
"""
from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def pt(value):
    return None if value is None else round(value.pt, 2)


def style_summary(style):
    font = style.font
    pf = style.paragraph_format
    xml = style._element.xml
    return {
        "name": style.name,
        "style_id": style.style_id,
        "type": str(style.type),
        "font_name": font.name,
        "font_size_pt": pt(font.size),
        "bold": font.bold,
        "line_spacing": str(pf.line_spacing),
        "space_before_pt": pt(pf.space_before),
        "space_after_pt": pt(pf.space_after),
        "first_line_indent_pt": pt(pf.first_line_indent),
        "left_indent_pt": pt(pf.left_indent),
        "has_numbering": "numPr" in xml,
        "outline_level": "outlineLvl" in xml,
    }


def paragraph_summary(paragraph, idx):
    runs = []
    for run in paragraph.runs[:8]:
        if not run.text:
            continue
        color = None
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
        runs.append(
            {
                "text": run.text[:80],
                "font": run.font.name,
                "size_pt": pt(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "color": color,
                "superscript": run.font.superscript,
            }
        )
    return {
        "index": idx,
        "style": paragraph.style.name,
        "text": paragraph.text[:200],
        "alignment": str(paragraph.alignment),
        "has_drawing": "<w:drawing" in paragraph._element.xml,
        "has_omath": "m:oMath" in paragraph._element.xml,
        "runs": runs,
    }


def table_summary(table, idx):
    preview = []
    for row in table.rows[:3]:
        preview.append([cell.text[:80] for cell in row.cells[:6]])
    return {
        "index": idx,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "preview": preview,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("template")
    parser.add_argument("--out", default=None)
    args = parser.parse_args()

    path = Path(args.template)
    doc = Document(path)

    report = {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "styles_used": sorted({p.style.name for p in doc.paragraphs}),
        "styles": [
            style_summary(style)
            for style in doc.styles
            if style.name in {p.style.name for p in doc.paragraphs}
        ],
        "paragraphs_first_120": [
            paragraph_summary(p, i) for i, p in enumerate(doc.paragraphs[:120])
        ],
        "heading_paragraphs": [
            paragraph_summary(p, i)
            for i, p in enumerate(doc.paragraphs)
            if p.style.name.startswith("Heading")
        ],
        "tables": [table_summary(t, i) for i, t in enumerate(doc.tables)],
    }

    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        if "word/document.xml" in names:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            report["document_xml_hints"] = {
                "hyperlink_count": xml.count("<w:hyperlink"),
                "bookmark_count": xml.count("<w:bookmarkStart"),
                "omath_count": xml.count("<m:oMath"),
                "blue_hyperlink_color": (
                    'w:color w:val="0563C1"' in xml
                    or 'w:color w:val="0000FF"' in xml
                ),
            }
        if "word/numbering.xml" in names:
            numbering = zf.read("word/numbering.xml").decode("utf-8", errors="ignore")
            report["numbering_xml_length"] = len(numbering)

    out = Path(args.out) if args.out else path.with_suffix(".template-report.json")
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
